"""
pdf_parser.py
=============
Extracts lab values from patient-uploaded documents.

Supported formats:
  1. Kanva PDF  (pdfplumber text extraction + targeted regex)
  2. Word .docx (python-docx text extraction + same regex)
  3. KIMS Srishti HTML (BeautifulSoup, line-pair parser — reused from v3_parser)
  4. Generic PDF fallback (pdfplumber, best-effort)

Returns:
  {
    'lab_values':    dict of {field: float},   # only present fields
    'demographics':  {'name', 'age', 'gender', 'reg_no', 'date'},
    'source_type':   str,   # 'kanva_pdf' | 'word' | 'kims_html' | 'generic_pdf'
    'n_labs_found':  int,
    'parse_warnings': list[str],
  }

Usage:
    from pdf_parser import parse_lab_report
    result = parse_lab_report('path/to/report.pdf')
    lab_values = result['lab_values']
"""

import re
import io
from pathlib import Path
from typing import Optional


# ══════════════════════════════════════════════════════════════════
#  MASTER TEST NAME → SCHEMA FIELD MAP
#  Covers Kanva PDF names, KIMS Srishti HTML names, Word doc names,
#  and common generic lab report variations.
# ══════════════════════════════════════════════════════════════════
TEST_MAP = {
    # ── Hemoglobin ─────────────────────────────────────────────────
    'HEMOGLOBIN': 'hemoglobin',
    'HAEMOGLOBIN': 'hemoglobin',
    'HB': 'hemoglobin',
    'HEMOGLOBIN (SLS HEMOGLOBIN(AUTOMATED))': 'hemoglobin',
    'HEMOGLOBIN (SLS)': 'hemoglobin',
    'HEMOGLOBIN (SLS HEMOGLOBIN(AUTOMATED))': 'hemoglobin',
    'HAEMOGLOBIN (SLS)': 'hemoglobin',
    'HAEMOGLOBIN (SLS HAEMOGLOBIN(AUTOMATED))': 'hemoglobin',
    # Kanva-specific
    'HAEMOGLOBIN(SLS)': 'hemoglobin',

    # ── Platelets ──────────────────────────────────────────────────
    'PLATELET COUNT': 'platelets',
    'PLATELET COUNT (AUTOMATED/MANUAL)': 'platelets',
    'PLATELETS': 'platelets',
    'PLATELET COUNT (DC DETECTION METHOD BY HYDRODYNAMIC FOCUSING)': 'platelets',
    'PLATELET COUNT(DC DETECTION METHOD BY HYDRODYNAMIC FOCUSING)': 'platelets',

    # ── WBC ────────────────────────────────────────────────────────
    'TOTAL COUNT': 'wbc',
    'TOTAL LEUCOCYTE COUNT': 'wbc',
    'TOTAL LEUCOCYTE COUNT (FLOW CYTOMETRY)': 'wbc',
    'TOTAL COUNT(TC)': 'wbc',
    'TC': 'wbc',
    'WBC': 'wbc',
    'WBC COUNT': 'wbc',
    'WBC COUNT (FLOURESCENT FLOWCYTOMETRY)': 'wbc',
    'TOTAL COUNT(TC) (METHOD-FLOW CYTOMETRY)': 'wbc',

    # ── ESR ────────────────────────────────────────────────────────
    'ESR': 'esr',
    'ESR (AUTOMATED)': 'esr',
    'ESR (WESTERGREN METHOD)': 'esr',

    # ── Bilirubin ──────────────────────────────────────────────────
    'TOTAL BILIRUBIN': 'bilirubin_total',
    'TOTAL BILIRUBIN(METHOD-JENDRASSIK - GROF)': 'bilirubin_total',
    'TOTAL BILIRUBIN(METHOD-JENDRASSIK-GROF)': 'bilirubin_total',

    # ── Albumin ────────────────────────────────────────────────────
    'SERUM ALBUMIN': 'albumin',
    'SERUM ALBUMIN(METHOD-BCG)': 'albumin',
    'ALBUMIN': 'albumin',
    'ALBUMIN(BCG)': 'albumin',
    # Kanva uses just "ALBUMIN" in LFT section — already covered above

    # ── Ferritin ───────────────────────────────────────────────────
    'FERRITIN': 'ferritin',
    'SERUM FERRITIN': 'ferritin',

    # ── Serum Iron ─────────────────────────────────────────────────
    'IRON': 'serum_iron',
    'SERUM IRON': 'serum_iron',
    'SERUM IRON(METHOD-FERROZINE)': 'serum_iron',

    # ── Creatinine ─────────────────────────────────────────────────
    'CREATININE': 'creatinine',
    'SERUM CREATININE': 'creatinine',
    'CREATININE (METHOD-JAFFES IDMS TRACEABLE)': 'creatinine',
    'CREATININE(METHOD-JAFFES IDMS TRACEABLE)': 'creatinine',
    'CREATININE(METHOD-JAFFE\'S KINETIC - IDMS)': 'creatinine',
    "CREATININE(JAFFE'S KINETIC - IDMS)": 'creatinine',

    # ── Urea ───────────────────────────────────────────────────────
    'UREA': 'urea',
    'BLOOD UREA': 'urea',
    'UREA (METHOD-GLDH/UREASE)': 'urea',
    'UREA(METHOD-GLDH/UREASE)': 'urea',

    # ── Glucose / GRBS ─────────────────────────────────────────────
    'GRBS': 'grbs',
    'GLUCOSE FASTING': 'grbs',
    'BLOOD GLUCOSE FASTING': 'grbs',
    'GLUCOSE RANDOM': 'grbs',
    'RANDOM BLOOD SUGAR': 'grbs',
    'FBS': 'grbs',
    'RBS': 'grbs',
    'GLUCOSE FASTING(HEXOKINASE)': 'grbs',

    # ── HbA1c ──────────────────────────────────────────────────────
    'GLYCOSYLATED HB (HBA1C)': 'hba1c',
    'GLYCOSYLATED HAEMOGLOBIN(HBA1C)': 'hba1c',
    'GLYCOSYLATED HAEMOGLOBIN(HBAIC)': 'hba1c',
    'HBA1C': 'hba1c',
    'GLYCOSYLATED HEMOGLOBIN': 'hba1c',
    'GLYCOSYLATED HAEMOGLOBIN - HBA1C': 'hba1c',
    'GLYCOSYLATED HAEMOGLOBIN(HBA1C)': 'hba1c',

    # ── TSAT ───────────────────────────────────────────────────────
    'TRANSFERRIN SATURATION': 'tsat',
    'TSAT': 'tsat',
    'TRANSFERRIN SATURATION(METHOD-CALCULATED)': 'tsat',

    # ── EF ─────────────────────────────────────────────────────────
    'EF': 'ef_percent',
    'EJECTION FRACTION': 'ef_percent',

    # ── BMI ────────────────────────────────────────────────────────
    'BMI': 'bmi',
    'BODY MASS INDEX': 'bmi',

    # ── Sodium ─────────────────────────────────────────────────────
    'SODIUM': 'sodium',
    'SODIUM(METHOD-ISE INDIRECT)': 'sodium',
    'SODIUM(ISE INDIRECT)': 'sodium',

    # ── Potassium ──────────────────────────────────────────────────
    'POTASSIUM': 'potassium',
    'POTASSIUM(METHOD-ISE INDIRECT)': 'potassium',
    'POTASSIUM(ISE INDIRECT)': 'potassium',

    # ── Chloride ───────────────────────────────────────────────────
    'CHLORIDE': 'chloride',
    'CHLORIDE(METHOD-ISE INDIRECT)': 'chloride',
    'CHLORIDE(ISE INDIRECT)': 'chloride',

    # ── SGOT ───────────────────────────────────────────────────────
    'SGOT': 'sgot',
    'SGOT(AST)': 'sgot',
    'SGOT(AST)(METHOD-IFCC WITH P5P)': 'sgot',
    'S.G.O.T (A.S.T)': 'sgot',
    'S.G.O.T(A.S.T)': 'sgot',
    'SGOT(AST)(UV WITHOUT P5P)': 'sgot',   # Kanva uses "UV without P5P"
    'SGOT (AST)(UV WITHOUT P5P)': 'sgot',

    # ── SGPT ───────────────────────────────────────────────────────
    'SGPT': 'sgpt',
    'SGPT(ALT)': 'sgpt',
    'SGPT(ALT)(METHOD-IFCC WITH P5P)': 'sgpt',
    'S.G.P.T (A.L.T)': 'sgpt',
    'S.G.P.T(A.L.T)': 'sgpt',
    'SGPT(ALT)(UV WITHOUT P5P)': 'sgpt',
    'SGPT (ALT)(UV WITHOUT P5P)': 'sgpt',

    # ── ALP ────────────────────────────────────────────────────────
    'ALKALINE PHOSPHATASE': 'alp',
    'ALKALINE PHOSPHATASE(METHOD-IFCC)': 'alp',
    'ALKALINE PHOSPHATASE(PNPP)': 'alp',

    # ── GGT ────────────────────────────────────────────────────────
    'GGT': 'ggt',
    'GAMMA GLUTAMYL TRASNSFERASE (GGT)': 'ggt',
    'GAMMA GLUTAMYL TRASNSFERASE (GGT)(METHOD-IFCC)': 'ggt',
    'GAMMA GT': 'ggt',
    'GAMMA-GT': 'ggt',
    'GAMMA GT(G-GLUATAMYL-CARBOXY-NITRONILIDE)': 'ggt',

    # ── CRP ────────────────────────────────────────────────────────
    'CRP': 'crp',
    'C-REACTIVE PROTEIN': 'crp',
    'C - REACTIVE PROTEIN': 'crp',
    'CRP(METHOD-IMMUNOTURBIDIMETRY)': 'crp',

    # ── Uric Acid ──────────────────────────────────────────────────
    'URIC ACID': 'uric_acid',
    'SERUM URIC ACID': 'uric_acid',
    'SERUM URIC ACID(METHOD-URICASE)': 'uric_acid',
    'URIC ACID(METHOD-URICASE)': 'uric_acid',
    'URIC ACID(URICASE, COLORIMETRIC)': 'uric_acid',

    # ── Neutrophils ────────────────────────────────────────────────
    'NEUTROPHILS': 'neutrophils_pct',
    'NEUTROPHILS (FLOW CYTOMETRY)': 'neutrophils_pct',
    'NEUTROPHILS (METHOD-FLOW CYTOMETRY)': 'neutrophils_pct',

    # ── Lymphocytes ────────────────────────────────────────────────
    'LYMPHOCYTES': 'lymphocytes_pct',
    'LYMPHOCYTES (FLOW CYTOMETRY)': 'lymphocytes_pct',
    'LYMPHOCYTES (METHOD-FLOW CYTOMETRY)': 'lymphocytes_pct',

    # ── MCV ────────────────────────────────────────────────────────
    'MCV': 'mcv',
    'MCV (CALCULATED)': 'mcv',
    'MCV(CALCULATED)': 'mcv',

    # ── RDW ────────────────────────────────────────────────────────
    'RDW-CV': 'rdw',
    'RDW-CV (AUTOMATED)': 'rdw',
    'RDW': 'rdw',
    'RDW (ELECTRICAL IMPEDANCE VARIATION /FLOURESCENT FLOWCYTOMETRY)': 'rdw',

    # ── PCV ────────────────────────────────────────────────────────
    'PACKED CELL VOLUME': 'pcv',
    'PACKED CELL VOLUME (CALCULATED)': 'pcv',
    'PACKED CELL VOLUME(PCV)': 'pcv',
    'PCV (HAEMATOCRIT) (CALCULATED)': 'pcv',
    'PCV(HAEMATOCRIT)(CALCULATED)': 'pcv',
}

# Fields to skip — not in schema, found in Kanva/KIMS reports
SKIP_FIELDS = {
    'DIRECT BILIRUBIN', 'INDIRECT BILIRUBIN', 'TOTAL PROTEIN',
    'GLOBULIN', 'AG RATIO', 'A/G RATIO', 'EAG',
    'ESTIMATED AVERAGE GLUCOSE', 'RBC COUNT', 'RED BLOOD CORPULSES',
    'RBC', 'MCH', 'MCHC', 'MONOCYTES', 'EOSINOPHILS', 'BASOPHILS',
    'T.CHOLESTEROL', 'TRIGLYCERIDES', 'HDL', 'LDL', 'VLDL',
    'TC/HDL RATIO', 'LDL/HDL RATIO', 'T3', 'T4', 'TSH',
    'VITAMIN D', 'VITAMIN B12', 'FOLATE', 'PSA',
    'GLUCOSE 2HR POST PRANDIAL', 'GLUCOSE 2 HR POST PRANDIAL',
    'PERIPHERAL SMEAR', 'IMPRESSION',
}

# Noise lines to skip during parsing
NOISE_PATTERNS = [
    r'^©', r'^Srishti', r'APPROVED', r'Lab Reports', r'New Search',
    r'Last Login', r'^User:', r'^Role:', r'^Station:', r'My Account',
    r'^Logout', r'END OF THE REPORT', r'End of Report',
    r'Prepared By', r'Verified By', r'APPROVED BY',
    r'Biological Reference', r'Reference Interval',
    r'Reported On', r'Sample Received', r'^\.$',
    r'^-+$', r'^=+$', r'DISCLAIMER', r'The result obtained',
    r'Page \d+ of \d+', r'www\.', r'E-Mail', r'WebSite',
    r'PH:', r'Tel\.', r'Reg\.No', r'KMC-',
]


def _normalize_name(name: str) -> str:
    """Normalize test name for dict lookup."""
    n = name.strip()
    # Remove method suffixes in parens for secondary lookup
    n_clean = re.sub(r'\s*\(Method[^)]*\)', '', n, flags=re.IGNORECASE)
    n_clean = re.sub(r'\s+', ' ', n_clean).strip().upper()
    return n_clean


def _parse_numeric(val_str: str, field: str) -> Optional[float]:
    """Extract numeric value from a result string, handling unit words."""
    if not val_str:
        return None
    val_str = str(val_str).strip()

    # Handle platelet lakhs format
    if field == 'platelets':
        m = re.search(r'([\d\.]+)', val_str)
        if m:
            v = float(m.group(1))
            if v < 100:   # value in lakhs
                return round(v * 100000)
            return round(v)
        return None

    # Standard numeric extraction
    m = re.search(r'([\d]+\.[\d]+|[\d]+)', val_str)
    if m:
        try:
            return float(m.group())
        except ValueError:
            return None
    return None


def _is_noise(line: str) -> bool:
    for pat in NOISE_PATTERNS:
        if re.search(pat, line, re.IGNORECASE):
            return True
    return False


# ══════════════════════════════════════════════════════════════════
#  KANVA PDF PARSER
# ══════════════════════════════════════════════════════════════════

def _parse_kanva_pdf(filepath: str) -> dict:
    """
    Parse a Kanva Diagnostic Services PDF.
    Structure per page: Test Parameter | Result(s) | Biological Reference Interval
    pdfplumber extracts this as text with consistent column layout.
    """
    try:
        import pdfplumber
    except ImportError:
        return {'error': 'pdfplumber not installed'}

    lab_values = {}
    demographics = {}
    warnings = []

    try:
        with pdfplumber.open(filepath) as pdf:
            full_text = ''
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if text:
                    full_text += text + '\n'
    except Exception as e:
        return {'error': f'pdfplumber failed: {e}'}

    lines = [l.strip() for l in full_text.split('\n') if l.strip()]

    # ── Extract patient demographics ───────────────────────────────
    for i, line in enumerate(lines[:30]):
        # Kanva header: "Reg. No. K3190105"
        m = re.search(r'Reg\.\s*No\.?\s+([A-Z0-9]+)', line)
        if m:
            demographics['reg_no'] = m.group(1)
        m = re.search(r'Name\s+(.+)', line)
        if m and 'name' not in demographics:
            demographics['name'] = m.group(1).strip()
        m = re.search(r'Age\s+(\d+)\s*Year', line, re.IGNORECASE)
        if m:
            demographics['age'] = int(m.group(1))
        m = re.search(r'Sex\s+(Male|Female)', line, re.IGNORECASE)
        if m:
            demographics['gender'] = 1 if m.group(1).lower() == 'male' else 0
        m = re.search(r'Date\s+(\d{2}/\d{2}/\d{4})', line)
        if m:
            demographics['date'] = m.group(1)

    # ── Parse test results ─────────────────────────────────────────
    # Kanva format after header:
    #   TEST NAME
    #   VALUE  UNIT
    #   reference range text...
    # OR same-line: "TEST NAME  VALUE  UNIT  REFERENCE"

    # Kanva inline pattern: "TEST NAME (method) VALUE UNIT reference..."
    # pdfplumber extracts each test inline: all on one line
    INLINE_PAT = re.compile(r'^(.+?)\s+([\d]+\.[\d]+|[\d]+)\s+([a-zA-Z%/µ]+)')

    SKIP_STARTS = (
        'CLINICAL BIOCHEMISTRY', 'LIVER FUNCTION TEST', 'HAEMATOLOGY',
        'SPECIAL BIOCHEMISTRY', 'URINE REPORT', 'ELECTROLYTES',
        'COMPLETE BLOOD COUNT', 'DIFFERENTIAL COUNT', 'RED CELL ABSOLUTE',
        'GLYCOSYLATED HAEMOGLOBIN -', 'T3,T4,TSH', 'PERIPHRAL SMEAR',
        'TEST PARAMETER', 'BIOLOGICAL REFERENCE', 'TOTAL & DIFFERENTIAL',
        'COMPLETE HAEMOGRAM', 'PERIPHERAL SMEAR', 'CHEMICAL EXAMINATION',
        'MICROSCOPIC EXAMINATION', 'PHYSICAL EXAMINATION',
        'SAMPLE RECEIVED', 'VERIFIED BY', 'REPORTED ON',
    )

    for line in lines:
        if not line or _is_noise(line):
            continue
        upper = line.strip().upper()
        if any(upper.startswith(s) for s in SKIP_STARTS):
            continue

        m = INLINE_PAT.match(line)
        if not m:
            continue

        test_name = m.group(1).strip()
        value_str = m.group(2)
        norm = _normalize_name(test_name)
        field = (TEST_MAP.get(norm)
                 or TEST_MAP.get(re.sub(r'\s*\([^)]*\)', '', norm).strip()))
        if field and field not in lab_values:
            val = _parse_numeric(value_str, field)
            if val is not None:
                lab_values[field] = val
        i += 1

    return {
        'lab_values':     lab_values,
        'demographics':   demographics,
        'source_type':    'kanva_pdf',
        'n_labs_found':   len(lab_values),
        'parse_warnings': warnings,
    }


# ══════════════════════════════════════════════════════════════════
#  GENERIC PDF PARSER (fallback for non-Kanva PDFs)
#  Also handles lab reports like Dhirghyau Hospital style
# ══════════════════════════════════════════════════════════════════

def _parse_generic_pdf(filepath: str) -> dict:
    """
    Fallback PDF parser: pdfplumber text extraction with flexible
    regex matching for NAME : VALUE patterns common across lab reports.
    """
    try:
        import pdfplumber
    except ImportError:
        return {'error': 'pdfplumber not installed'}

    lab_values = {}
    demographics = {}
    warnings = []

    try:
        with pdfplumber.open(filepath) as pdf:
            full_text = ''
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text += text + '\n'
    except Exception as e:
        return {'error': f'PDF read failed: {e}'}

    lines = [l.strip() for l in full_text.split('\n') if l.strip()]

    # ── Demographics ───────────────────────────────────────────────
    for line in lines[:40]:
        m = re.search(r'AGE\s*[:\-]?\s*(\d+)', line, re.IGNORECASE)
        if m:
            demographics['age'] = int(m.group(1))
        m = re.search(r'SEX\s*[:\-]?\s*(Male|Female)', line, re.IGNORECASE)
        if m:
            demographics['gender'] = 1 if m.group(1).lower() == 'male' else 0
        m = re.search(r'NAME\s*[:\-]\s*(.+)', line, re.IGNORECASE)
        if m and 'name' not in demographics:
            demographics['name'] = m.group(1).strip()

    # ── Line-by-line value extraction ─────────────────────────────
    num_pat = re.compile(r'^[\d\.\+\-<>]+$')

    i = 0
    while i < len(lines):
        line = lines[i]
        if _is_noise(line):
            i += 1
            continue

        # Pattern 1: "TEST NAME : VALUE UNIT"
        m = re.match(r'^(.+?)\s*[:\-]\s*([\d\.]+)\s*', line)
        if m:
            test_name = m.group(1).strip()
            value_str = m.group(2)
            norm = _normalize_name(test_name)
            field = TEST_MAP.get(norm) or TEST_MAP.get(
                re.sub(r'\s*\([^)]*\)', '', norm).strip()
            )
            if field and field not in lab_values:
                val = _parse_numeric(value_str, field)
                if val is not None:
                    lab_values[field] = val
            i += 1
            continue

        # Pattern 2: line-pair (test name on one line, value on next)
        if i + 1 < len(lines) and num_pat.match(lines[i + 1].replace('.', '', 1)):
            norm = _normalize_name(line)
            field = TEST_MAP.get(norm) or TEST_MAP.get(
                re.sub(r'\s*\([^)]*\)', '', norm).strip()
            )
            if field and field not in lab_values:
                val = _parse_numeric(lines[i + 1], field)
                if val is not None:
                    lab_values[field] = val
            i += 2
            continue

        i += 1

    if not lab_values:
        warnings.append('Generic PDF parser found no lab values — report format unrecognised')

    return {
        'lab_values':     lab_values,
        'demographics':   demographics,
        'source_type':    'generic_pdf',
        'n_labs_found':   len(lab_values),
        'parse_warnings': warnings,
    }


# ══════════════════════════════════════════════════════════════════
#  WORD (.docx) PARSER
# ══════════════════════════════════════════════════════════════════

def _parse_word_doc(filepath: str) -> dict:
    """
    Parse a Word .docx lab report.
    Extracts all text, then runs same line-pair + colon-delimited regex.
    """
    try:
        import docx
    except ImportError:
        return {'error': 'python-docx not installed'}

    lab_values = {}
    demographics = {}
    warnings = []

    try:
        doc = docx.Document(filepath)
        lines = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                lines.append(t)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    lines.append(row_text)
    except Exception as e:
        return {'error': f'python-docx failed: {e}'}

    num_pat = re.compile(r'^[\d\.\+\-<>]+$')

    for line in lines[:30]:
        m = re.search(r'AGE\s*[:\-]?\s*(\d+)', line, re.IGNORECASE)
        if m:
            demographics['age'] = int(m.group(1))
        m = re.search(r'SEX\s*[:\-]?\s*(Male|Female)', line, re.IGNORECASE)
        if m:
            demographics['gender'] = 1 if m.group(1).lower() == 'male' else 0
        m = re.search(r'NAME\s*[:\-]\s*(.+)', line, re.IGNORECASE)
        if m and 'name' not in demographics:
            demographics['name'] = m.group(1).strip()

    i = 0
    while i < len(lines):
        line = lines[i]
        if _is_noise(line):
            i += 1
            continue

        # Table row: "TEST NAME | VALUE | UNIT | RANGE"
        if '|' in line:
            parts = [p.strip() for p in line.split('|')]
            if len(parts) >= 2:
                test_name = parts[0]
                value_str = parts[1]
                norm = _normalize_name(test_name)
                field = TEST_MAP.get(norm) or TEST_MAP.get(
                    re.sub(r'\s*\([^)]*\)', '', norm).strip()
                )
                if field and field not in lab_values:
                    val = _parse_numeric(value_str, field)
                    if val is not None:
                        lab_values[field] = val
            i += 1
            continue

        # Colon pattern
        m = re.match(r'^(.+?)\s*[:\-]\s*([\d\.]+)', line)
        if m:
            norm = _normalize_name(m.group(1))
            field = TEST_MAP.get(norm) or TEST_MAP.get(
                re.sub(r'\s*\([^)]*\)', '', norm).strip()
            )
            if field and field not in lab_values:
                val = _parse_numeric(m.group(2), field)
                if val is not None:
                    lab_values[field] = val
            i += 1
            continue

        # Line-pair
        if i + 1 < len(lines) and num_pat.match(lines[i + 1].replace('.', '', 1)):
            norm = _normalize_name(line)
            field = TEST_MAP.get(norm) or TEST_MAP.get(
                re.sub(r'\s*\([^)]*\)', '', norm).strip()
            )
            if field and field not in lab_values:
                val = _parse_numeric(lines[i + 1], field)
                if val is not None:
                    lab_values[field] = val
            i += 2
            continue

        i += 1

    return {
        'lab_values':     lab_values,
        'demographics':   demographics,
        'source_type':    'word_doc',
        'n_labs_found':   len(lab_values),
        'parse_warnings': warnings,
    }


# ══════════════════════════════════════════════════════════════════
#  KIMS SRISHTI HTML PARSER (adapted from v3_parser.py)
# ══════════════════════════════════════════════════════════════════

def _parse_kims_html(filepath: str) -> dict:
    """
    Parse a single KIMS Srishti LIS HTML report.
    Line-pair structure: test name line → value line → unit line → ref line.
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return {'error': 'beautifulsoup4 not installed'}

    lab_values = {}
    demographics = {}

    try:
        content = open(filepath, encoding='utf-8', errors='ignore').read()
        soup = BeautifulSoup(content, 'html.parser')
        text = soup.get_text(separator='\n', strip=True)
    except Exception as e:
        return {'error': f'HTML parse failed: {e}'}

    lines = [l.strip() for l in text.split('\n') if l.strip()]

    # ── Demographics ───────────────────────────────────────────────
    for i, line in enumerate(lines[:35]):
        if 'Age/Gender' in line and i + 1 < len(lines):
            ag = lines[i + 1].lstrip(':').strip()
            m = re.search(r'(\d+)Y\s*/\s*(Male|Female)', ag, re.IGNORECASE)
            if m:
                demographics['age'] = int(m.group(1))
                demographics['gender'] = 1 if m.group(2).lower() == 'male' else 0
        if 'Patient Name' in line and i + 1 < len(lines):
            demographics['name'] = lines[i + 1].lstrip(':').strip()
        if 'UHID' in line and i + 1 < len(lines):
            demographics['reg_no'] = lines[i + 1].lstrip(':').strip()

    # ── Find test section ──────────────────────────────────────────
    start = 0
    for i, line in enumerate(lines):
        if 'Biological Reference Interval' in line or 'Test Name' in line:
            start = i + 1
            break

    num_pat = re.compile(r'^[\d\.\+\-<>]+$')
    noise_kw = {
        '© 2026', 'Srishti', 'APPROVED', 'Lab Reports', 'New Search',
        'Last Login', 'User:', 'Role:', 'Station:', 'My Account',
        'Logout', '2026-2027', '2025-2026', 'END OF THE REPORT',
        'Prepared By', 'Verified By', 'APPROVED BY', '.',
        'Biological Reference Interval', 'Test Name', 'Results',
        'Unit', 'DEPARTMENT OF',
    }

    i = start
    while i < len(lines):
        line = lines[i]
        if any(n in line for n in noise_kw):
            i += 1
            continue
        if 'END OF THE REPORT' in line or 'Prepared By' in line:
            break

        if i + 1 < len(lines) and num_pat.match(lines[i + 1].replace('.', '', 1)):
            raw_name = line
            raw_val  = lines[i + 1]
            norm = _normalize_name(raw_name)
            field = TEST_MAP.get(norm) or TEST_MAP.get(
                re.sub(r'\s*\([^)]*\)', '', norm).strip()
            )
            if field and field not in lab_values:
                val = _parse_numeric(raw_val, field)
                if val is not None:
                    lab_values[field] = val
            i += 2
        else:
            i += 1

    return {
        'lab_values':     lab_values,
        'demographics':   demographics,
        'source_type':    'kims_html',
        'n_labs_found':   len(lab_values),
        'parse_warnings': [],
    }


# ══════════════════════════════════════════════════════════════════
#  PUBLIC ENTRY POINT
# ══════════════════════════════════════════════════════════════════

def parse_lab_report(filepath: str) -> dict:
    """
    Auto-detect format and parse a lab report file.

    Args:
        filepath: path to .pdf, .docx, or .html file

    Returns:
        {
          'lab_values':     dict,     # {field: float} — only present fields
          'demographics':   dict,     # {name, age, gender, reg_no, date}
          'source_type':    str,
          'n_labs_found':   int,
          'parse_warnings': list[str],
        }

    On parse failure, lab_values will be empty and parse_warnings will explain.
    """
    path = Path(filepath)
    ext  = path.suffix.lower()

    if ext == '.html' or ext == '.htm':
        result = _parse_kims_html(filepath)

    elif ext == '.docx':
        result = _parse_word_doc(filepath)

    elif ext == '.pdf':
        # Try Kanva parser first (more structured), fall back to generic
        result = _parse_kanva_pdf(filepath)

        # If Kanva got fewer than 3 labs, try generic as well and take the better result
        if result.get('n_labs_found', 0) < 3:
            generic = _parse_generic_pdf(filepath)
            if generic.get('n_labs_found', 0) > result.get('n_labs_found', 0):
                result = generic
                result['parse_warnings'].insert(
                    0, 'Kanva parser found <3 labs; fell back to generic PDF parser'
                )

    else:
        return {
            'lab_values':     {},
            'demographics':   {},
            'source_type':    'unsupported',
            'n_labs_found':   0,
            'parse_warnings': [f'Unsupported file type: {ext}'],
        }

    # Ensure 'error' key from sub-parsers is surfaced cleanly
    if 'error' in result:
        result.setdefault('lab_values', {})
        result.setdefault('demographics', {})
        result.setdefault('n_labs_found', 0)
        result['parse_warnings'] = result.get('parse_warnings', []) + [result.pop('error')]

    return result


# ══════════════════════════════════════════════════════════════════
#  CLI QUICK TEST
# ══════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    import sys, json
    if len(sys.argv) < 2:
        print('Usage: python pdf_parser.py <path/to/report.pdf|.docx|.html>')
        sys.exit(1)

    result = parse_lab_report(sys.argv[1])
    print(f"\nSource type  : {result['source_type']}")
    print(f"Labs found   : {result['n_labs_found']}")
    print(f"Demographics : {result['demographics']}")
    print(f"Warnings     : {result['parse_warnings']}")
    print('\nLab values:')
    for k, v in sorted(result['lab_values'].items()):
        print(f'  {k:<25} {v}')
